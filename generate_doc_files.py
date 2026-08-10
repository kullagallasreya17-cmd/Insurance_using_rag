from pathlib import Path
from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas

root = Path(r'd:\projects\Insurance_using_RAG')
content = '''
Enterprise Insurance RAG Application
==================================

Overview
--------
This project is an insurance-focused RAG application built with React and FastAPI. It supports document upload, semantic search over policy documents, grounded chatbot responses, claim analysis, and admin-style operations.

Key Features
------------
- PDF and image upload support
- Text extraction and chunking
- Embedding generation and vector storage with MongoDB Atlas Vector Search
- Grounded chat responses using retrieved context
- Claim analysis workflow
- Authentication and role-based access

Project Structure
-----------------
- backend/: FastAPI backend, auth, document processing, RAG logic, database models
- frontend/: React/Vite frontend UI for portal and chatbot
- services/: modular service layout for gateway, authentication, documents, AI, and RAG

Run Backend
-----------
cd D:\\projects\\Insurance_using_RAG\\backend
.\\.venv\\Scripts\\python.exe -m uvicorn main:app --reload --host 127.0.0.1 --port 8000

Run Frontend
------------
cd D:\\projects\\Insurance_using_RAG\\frontend
npm.cmd run dev

Default Admin
-------------
Username: admin
Password: admin123

Architecture Summary
-------------------
The application uses:
- React frontend
- FastAPI backend
- MongoDB Atlas Vector Search for vector storage
- Google Gemini embeddings
- JWT-based authentication
- Document ingestion and retrieval pipeline

Notes
-----
This document was generated for the current workspace and can be used as a project handover or reference file.
'''

# Create DOCX
word_doc = Document()
word_doc.add_heading('Enterprise Insurance RAG Application', level=1)
for paragraph in content.splitlines():
    if not paragraph.strip():
        continue
    if paragraph.startswith('---'):
        continue
    if paragraph.startswith('Overview') or paragraph.startswith('Key Features') or paragraph.startswith('Project Structure') or paragraph.startswith('Run Backend') or paragraph.startswith('Run Frontend') or paragraph.startswith('Default Admin') or paragraph.startswith('Architecture Summary') or paragraph.startswith('Notes'):
        word_doc.add_heading(paragraph, level=2)
    else:
        word_doc.add_paragraph(paragraph)
word_doc.save(root / 'PROJECT_DOCUMENTATION.docx')

# Create PDF
pdf_path = root / 'PROJECT_DOCUMENTATION.pdf'
canvas_obj = canvas.Canvas(str(pdf_path), pagesize=letter)
text = content.replace('\\n', '\n')
lines = text.splitlines()
textobject = canvas_obj.beginText(40, 760)
for line in lines:
    textobject.textLine(line)
canvas_obj.drawText(textobject)
canvas_obj.save()

print('Generated documentation files:')
print(root / 'PROJECT_DOCUMENTATION.docx')
print(root / 'PROJECT_DOCUMENTATION.pdf')
